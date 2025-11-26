import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent.resolve()))
sys.path.append(str((Path(__file__).parent / "rag_pipeline").resolve()))
sys.path.append(str((Path(__file__).parent / "app_skeleton").resolve()))

import streamlit as st
from typing import List, Dict, Any, Tuple
from llama_index.core import Document

from app_skeleton.src.utils import get_file_type, format_response
from rag_pipeline.config.settings import RAGConfig
from rag_pipeline.core.pipeline import OptimizedRAGPipeline
from rag_pipeline.core.chat_pipeline import ChatPipeline
from rag_pipeline.core.ocr_extractor import PDFTextExtractor
from PIL import Image
import pytesseract



# Configurazione pagina
st.set_page_config(
    page_title="RAG Prefettura - Assistente per Illeciti Amministrativi",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personalizzato per migliorare l'aspetto
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f4e79;
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem;
        background: linear-gradient(90deg, #f0f8ff, #e6f3ff);
        border-radius: 10px;
        border-left: 5px solid #1f4e79;
    }
    
    .main-header-chat {
        font-size: 2.5rem;
        font-weight: 700;
        color: #7b1fa2;
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem;
        background: linear-gradient(90deg, #f3e5f5, #e1bee7);
        border-radius: 10px;
        border-left: 5px solid #7b1fa2;
    }
    
    .mode-toggle {
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem;
        background-color: #f8f9fa;
        border-radius: 10px;
        border: 2px solid #dee2e6;
    }
    
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 4px solid;
        color: #000000 !important;
    }
    
    .user-message {
        background-color: #e3f2fd;
        border-left-color: #2196f3;
    }
    
    .user-message-chat {
        background-color: #f3e5f5;
        border-left-color: #9c27b0;
    }
    
    .assistant-message {
        background-color: #f1f8e9;
        border-left-color: #4caf50;
    }
    
    .assistant-message-chat {
        background-color: #fce4ec;
        border-left-color: #e91e63;
    }
    
    .status-box {
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
        font-weight: bold;
    }
    
    .status-success {
        background-color: #d4edda;
        color: #155724;
        border: 1px solid #c3e6cb;
    }
    
    .status-error {
        background-color: #f8d7da;
        color: #721c24;
        border: 1px solid #f5c6cb;
    }
    
    .status-warning {
        background-color: #fff3cd;
        color: #856404;
        border: 1px solid #ffeaa7;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# 🆕 NUOVE FUNZIONI PER DETECTION ED EXTRACTION
# ============================================================================

def detect_extraction_method(file_path: str, file_type: str, force_ocr: bool) -> str:
    """
    Determina il metodo di estrazione basato su flag OCR e tipo file
    
    Args:
        file_path: Path del file
        file_type: Tipo file (extension senza punto)
        force_ocr: Flag da checkbox UI
    
    Returns:
        'ocr' | 'pdf_text' | 'docx' | 'text'
    """
    if force_ocr:
        # Se OCR forzato, usa OCR per tutto tranne TXT/MD
        if file_type in ['pdf', 'jpg', 'jpeg', 'png']:
            return 'ocr'
        elif file_type in ['docx', 'doc']:
            return 'docx'
        else:
            return 'text'
    else:
        # Detection automatica
        if file_type == 'pdf':
            return 'pdf_text'  # Prova testo prima, fallback OCR
        elif file_type in ['docx', 'doc']:
            return 'docx'
        elif file_type in ['jpg', 'jpeg', 'png']:
            return 'ocr'  # Immagini sempre OCR
        elif file_type in ['txt', 'md', 'text', 'markdown']:
            return 'text'
        else:
            raise ValueError(f"Tipo file non supportato: {file_type}")

def is_text_sufficient(text: str, min_chars: int = 100) -> bool:
    """
    Verifica se il testo estratto è sufficiente
    
    Args:
        text: Testo da verificare
        min_chars: Soglia minima caratteri
    
    Returns:
        True se testo sufficiente
    """
    return len(text.strip()) > min_chars

def extract_text_from_pdf(file_path: str) -> str:
    """Estrae testo da PDF usando fitz (PyMuPDF) direttamente"""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            text += f"\n--- Pagina {page_num + 1} ---\n"
            text += page_text
        
        doc.close()
        return text.strip()
        
    except ImportError:
        st.error("❌ PyMuPDF (fitz) non installato. Installa con: pip install pymupdf")
        return ""
    except Exception as e:
        st.warning(f"⚠️ Errore estrazione PDF: {e}")
        return ""

def extract_with_ocr(file_path: str) -> str:
    """
    Estrae testo usando OCR (Tesseract) con preprocessing avanzato
    Usa PDFTextExtractor per gestire PDF e immagini
    
    Args:
        file_path: Path del file
    
    Returns:
        Testo estratto con OCR
    """
    try:
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            # Usa PDFTextExtractor per PDF
            # Crea directory temporanea per output (richiesta dalla classe)
            temp_output = Path("data/temp_output")
            temp_output.mkdir(parents=True, exist_ok=True)
            
            # Instanzia extractor
            extractor = PDFTextExtractor(
                input_dir=str(Path(file_path).parent),
                output_dir=str(temp_output),
                dpi=300,
                ocr_threshold=50
            )
            
            # Estrai con OCR forzato
            pages_text = extractor.extract_text_from_pdf(
                file_path, 
                force_ocr=True  # Forza OCR per tutte le pagine
            )
            
            # Combina tutte le pagine
            text = ""
            for page_num in sorted(pages_text.keys()):
                text += f"\n--- Pagina {page_num} ---\n"
                text += pages_text[page_num]
            
            return text.strip()
        
        else:
            # Immagini (JPG, PNG) - usa preprocessing della classe
            temp_output = Path("data/temp_output")
            temp_output.mkdir(parents=True, exist_ok=True)
            
            extractor = PDFTextExtractor(
                input_dir=str(Path(file_path).parent),
                output_dir=str(temp_output),
                dpi=300,
                ocr_threshold=50
            )
            
            # Carica immagine
            img = Image.open(file_path)
            
            # Usa preprocessing della classe
            img = extractor.preprocess_image(img)
            
            # OCR con Tesseract
            text = pytesseract.image_to_string(img, lang='ita')
            return text.strip()
    
    except ImportError as e:
        st.error(f"❌ Dipendenze OCR mancanti: {e}")
        st.info("Installa con: pip install pytesseract opencv-python pillow pymupdf")
        return ""
    except Exception as e:
        st.warning(f"⚠️ Errore OCR: {e}")
        import traceback
        st.error(traceback.format_exc())
        return ""

def extract_from_docx(file_path: str) -> str:
    """
    Estrae testo da file DOCX
    
    Args:
        file_path: Path del DOCX
    
    Returns:
        Testo estratto
    """
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    
    except ImportError:
        st.error("❌ python-docx non installato. Installa con: pip install python-docx")
        return ""
    except Exception as e:
        st.warning(f"⚠️ Errore estrazione DOCX: {e}")
        return ""

def read_text_file(file_path: str) -> str:
    """
    Legge file di testo semplice
    
    Args:
        file_path: Path del file
    
    Returns:
        Contenuto del file
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        st.warning(f"⚠️ Errore lettura file: {e}")
        return ""

def extract_and_create_document(file_info: Dict[str, Any], extraction_method: str) -> Tuple[Document, str]:
    """
    Estrae testo e crea Document LlamaIndex con metadata completi
    
    Args:
        file_info: Dict con {filename, file_path, file_type, size}
        extraction_method: Metodo di estrazione da usare
    
    Returns:
        (Document, actual_method) dove actual_method può essere diverso se c'è fallback
    """
    file_path = file_info['file_path']
    filename = file_info['filename']
    file_type = file_info['file_type']
    
    text = ""
    actual_method = extraction_method
    
    # Estrazione basata sul metodo
    if extraction_method == 'pdf_text':
        # Tenta estrazione testo prima
        text = extract_text_from_pdf(file_path)
        
        # Smart fallback su OCR se testo insufficiente
        if not is_text_sufficient(text):
            st.warning(f"⚠️ {filename}: testo insufficiente, provo OCR...")
            text = extract_with_ocr(file_path)
            actual_method = 'pdf_ocr_fallback'
        else:
            actual_method = 'pdf_text'
    
    elif extraction_method == 'ocr':
        text = extract_with_ocr(file_path)
        actual_method = 'ocr'
    
    elif extraction_method == 'docx':
        text = extract_from_docx(file_path)
        actual_method = 'docx'
    
    elif extraction_method == 'text':
        text = read_text_file(file_path)
        actual_method = 'text'
    
    else:
        raise ValueError(f"Metodo di estrazione non valido: {extraction_method}")
    
    # Verifica che ci sia testo
    if not text or len(text.strip()) < 10:
        raise ValueError(f"Testo estratto insufficiente (< 10 caratteri)")
    
    # Crea Document LlamaIndex con metadata completi
    doc = Document(
        text=text,
        metadata={
            "source": file_path,
            "filename": filename,
            "file_type": file_type,
            "extraction_method": actual_method,
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": len(text.splitlines()),
            "file_size": file_info['size']
        }
    )
    
    return doc, actual_method

# ============================================================================
# 🔄 MODIFICATA: handle_file_upload con nuova logica
# ============================================================================

def handle_file_upload(uploaded_files, force_ocr: bool = False, batch_size: int = 50):
    """
    Gestisce upload con detection ed extraction automatica
    
    Args:
        uploaded_files: File caricati da st.file_uploader
        force_ocr: Flag per forzare OCR (da checkbox)
        batch_size: Dimensione batch per indicizzazione
    """
    if not uploaded_files:
        return
    
    successful_uploads = []
    failed_uploads = []
    
    upload_dir = Path("data/uploads")
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    # STEP 1: Salva file su disco
    for uploaded_file in uploaded_files:
        try:
            file_path = upload_dir / uploaded_file.name
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            file_type = get_file_type(uploaded_file.name)
            
            successful_uploads.append({
                'filename': uploaded_file.name,
                'file_type': file_type,
                'file_path': str(file_path),
                'size': uploaded_file.size
            })
                
        except Exception as e:
            failed_uploads.append({
                'filename': uploaded_file.name,
                'error': str(e)
            })
    
    if not successful_uploads:
        if failed_uploads:
            st.error("❌ Nessun file caricato con successo")
        return
    
    # STEP 2: Estrai testo e crea Documents
    new_documents = []
    extraction_info = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, file_info in enumerate(successful_uploads):
        try:
            status_text.text(f"🔍 Elaborando {file_info['filename']}...")
            
            # 2a. Determina metodo estrazione
            method = detect_extraction_method(
                file_info['file_path'],
                file_info['file_type'],
                force_ocr
            )
            
            # 2b. Estrai testo e crea Document
            doc, actual_method = extract_and_create_document(file_info, method)
            
            new_documents.append(doc)
            
            # Aggiorna file_info con metodo effettivo
            file_info['extraction_method'] = actual_method
            extraction_info.append(file_info)
            
        except Exception as e:
            st.error(f"❌ Errore elaborazione {file_info['filename']}: {e}")
            failed_uploads.append({
                'filename': file_info['filename'],
                'error': str(e)
            })
            continue
        
        finally:
            progress_bar.progress((i + 1) / len(successful_uploads))
    
    progress_bar.empty()
    status_text.empty()
    
    if not new_documents:
        st.error("❌ Nessun documento elaborato con successo")
        return
    
    # STEP 3: Accumula con documenti esistenti
    if 'all_documents' not in st.session_state:
        st.session_state.all_documents = []
    
    st.session_state.all_documents.extend(new_documents)
    
    # Aggiorna lista file info
    st.session_state.uploaded_files_info.extend(extraction_info)
    
    # STEP 4: Build index con TUTTI i documenti
    try:
        with st.spinner(f"📚 Indicizzando {len(st.session_state.all_documents)} documenti totali..."):
            # ✅ PASSA DOCUMENTS E BATCH_SIZE
            st.session_state.rag_pipeline.build_index(
                documents=st.session_state.all_documents,
                batch_size=batch_size  # 🆕 Parametro batch
            )
            
            # Setup query engine
            st.session_state.rag_pipeline.setup_query_engine()
            
            # Flag che l'indice è pronto
            st.session_state.index_ready = True
        
        st.success(f"✅ {len(new_documents)} nuovi documenti indicizzati! (Totale: {len(st.session_state.all_documents)})")
        
    except Exception as e:
        st.error(f"❌ Errore indicizzazione: {e}")
        import traceback
        st.error(traceback.format_exc())
        return
    
    # STEP 5: Mostra risultati
    if extraction_info:
        st.markdown('<div class="status-box status-success">✅ Documenti elaborati:</div>', unsafe_allow_html=True)
        for doc in extraction_info:
            # Icone basate sul metodo di estrazione
            icon = {
                'pdf_text': '📄',
                'pdf_ocr_fallback': '🔄',
                'ocr': '🔍',
                'docx': '📝',
                'text': '📃'
            }.get(doc['extraction_method'], '📄')
            
            st.write(f"{icon} {doc['filename']} ({doc['extraction_method']})")
    
    if failed_uploads:
        st.markdown('<div class="status-box status-error">❌ Errori:</div>', unsafe_allow_html=True)
        for doc in failed_uploads:
            st.write(f"• {doc['filename']}: {doc['error']}")

# ============================================================================
# FUNZIONI ESISTENTI (invariate)
# ============================================================================

def handle_query(query: str):
    """Gestisce le query dell'utente in base alla modalità"""
    if not query.strip():
        return
    
    current_mode = st.session_state.app_mode
    
    # Aggiungi messaggio utente alla cronologia corretta
    messages_key = f"messages_{current_mode}"
    st.session_state[messages_key].append({
        "role": "user",
        "content": query,
        "timestamp": ""
    })
    
    try:
        if current_mode == "rag":
            # Modalità RAG - verifica documenti caricati
            if not st.session_state.uploaded_files_info:
                st.error("⚠️ Carica almeno un documento prima di fare domande!")
                return
            
            if not st.session_state.get('index_ready', False):
                st.error("⚠️ Indice non costruito. Ricarica i documenti.")
                return
            
            with st.spinner("🤔 Elaborando la tua domanda con RAG..."):
                # Query RAG
                result = st.session_state.rag_pipeline.query(
                    query, 
                    enhance_query=True
                )
                
                response = result.get('answer', 'Nessuna risposta generata')
                sources = result.get('sources', [])
                metadata = result.get('query_metadata', {})
                
                # Aggiungi risposta con metadata completi
                st.session_state[messages_key].append({
                    "role": "assistant",
                    "content": response,
                    "sources": sources,
                    "metadata": metadata,
                    "response_time": result.get('response_time', 0),
                    "timestamp": ""
                })
        
        else:
            # ✅ Modalità Chat - PASSA CRONOLOGIA
            with st.spinner("💭 Pensando..."):
                if st.session_state.chat_pipeline:
                    # ✅ COSTRUISCI CRONOLOGIA (escludendo ultimo messaggio utente)
                    conversation_history = st.session_state[messages_key][:-1]
                    
                    # Usa ChatPipeline CON cronologia
                    result = st.session_state.chat_pipeline.query(
                        query,
                        conversation_history=conversation_history
                    )
                    
                    response = result.get('answer', 'Nessuna risposta generata')
                    
                    # ✅ AGGIUNGI METADATA CHAT
                    st.session_state[messages_key].append({
                        "role": "assistant",
                        "content": response,
                        "response_time": result.get('response_time', 0),
                        "model": result.get('model', ''),
                        "temperature": result.get('temperature', 0),
                        "timestamp": ""
                    })
                else:
                    # Fallback
                    response = "⚠️ ChatPipeline non inizializzato correttamente."
                    st.session_state[messages_key].append({
                        "role": "assistant",
                        "content": response,
                        "timestamp": ""
                    })
    
    except Exception as e:
        st.error(f"❌ Errore: {str(e)}")
        st.session_state[messages_key].append({
            "role": "assistant",
            "content": f"Mi dispiace, si è verificato un errore: {str(e)}",
            "timestamp": ""
        })

def display_chat_history():
    """Mostra la cronologia della chat in base alla modalità"""
    current_mode = st.session_state.app_mode
    messages_key = f"messages_{current_mode}"
    messages = st.session_state.get(messages_key, [])
    
    for message in messages:
        if message["role"] == "user":
            user_class = "user-message-chat" if current_mode == "chat" else "user-message"
            st.markdown(f"""
            <div class="chat-message {user_class}">
                <strong>🧑 Tu:</strong> {message["content"]}
            </div>
            """, unsafe_allow_html=True)
        else:
            assistant_class = "assistant-message-chat" if current_mode == "chat" else "assistant-message"
            icon = "💬" if current_mode == "chat" else "🤖"
            st.markdown(f"""
            <div class="chat-message {assistant_class}">
                <strong>{icon} Assistente:</strong> {format_response(message["content"])}
            </div>
            """, unsafe_allow_html=True)
            
            # ✅ MOSTRA METADATA ANCHE IN CHAT (ma più semplici)
            if current_mode == "chat" and "response_time" in message:
                with st.expander("ℹ️ Info risposta"):
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("⚡ Tempo", f"{message['response_time']:.2f}s")
                    with col2:
                        if "model" in message:
                            st.caption(f"🤖 {message['model']}")
                    with col3:
                        if "temperature" in message:
                            st.caption(f"🌡️ Temp: {message['temperature']}")
            
            # Mostra metadata RAG completi
            elif current_mode == "rag" and "metadata" in message and message["metadata"]:
                metadata = message["metadata"]
                
                with st.expander("🔍 Query Analysis & Pipeline Details"):
                    # Query Analysis
                    if 'expansions' in metadata:
                        st.markdown("**🔍 Query Analysis:**")
                        exp = metadata['expansions']
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            if 'intent' in exp:
                                st.info(f"**Intent:** {exp['intent']}")
                            if 'keywords' in exp and exp['keywords']:
                                keywords = ', '.join(exp['keywords'][:5])
                                st.info(f"**Keywords:** {keywords}")
                        
                        with col2:
                            if 'semantic_variants' in exp:
                                st.info(f"**Semantic variants:** {len(exp['semantic_variants'])}")
                            if 'num_queries_generated' in metadata:
                                st.info(f"**Total queries:** {metadata['num_queries_generated']}")
                    
                    st.divider()
                    
                    # Retrieval Pipeline
                    st.markdown("**📊 Retrieval Pipeline:**")
                    
                    if 'retrieval' in metadata:
                        ret = metadata['retrieval']
                        nodes = ret.get('total_nodes_retrieved', 0)
                        st.success(f"**1. Multi-retrieval:** {nodes} nodes retrieved")
                    
                    if 'deduplication' in metadata:
                        dedup = metadata['deduplication']
                        before = dedup.get('nodes_before', 0)
                        after = dedup.get('nodes_after', 0)
                        removed = dedup.get('duplicates_removed', 0)
                        st.success(f"**2. Deduplication:** {before} → {after} nodes (-{removed} duplicates)")
                    
                    if 'reranking' in metadata:
                        rerank = metadata['reranking']
                        if rerank.get('applied'):
                            before = rerank.get('nodes_before', 0)
                            after = rerank.get('nodes_after', 0)
                            st.success(f"**3. Reranking:** {before} → {after} top nodes ✓")
                        else:
                            st.warning("**3. Reranking:** Not applied")
                    
                    if 'response_time' in message:
                        st.divider()
                        st.metric("⚡ Total time", f"{message['response_time']:.3f}s")
            
            # Mostra fonti SOLO in modalità RAG
            if current_mode == "rag" and "sources" in message and message["sources"]:
                with st.expander(f"📚 Top Sources ({len(message['sources'])} total)"):
                    for i, source in enumerate(message['sources'][:3], 1):
                        st.markdown(f"**[{i}] Score: {source['score']:.3f}**")
                        st.text(source['text'][:200] + "...")
                        
                        if source.get('metadata'):
                            st.caption(f"📄 {source['metadata'].get('file_name', 'N/A')}")
                        
                        if source.get('reranked'):
                            st.success("✓ Reranked")
                        
                        st.divider()

def initialize_session_state():
    """Inizializza lo stato della sessione"""
    if 'app_mode' not in st.session_state:
        st.session_state.app_mode = 'rag'
    
    if 'messages_rag' not in st.session_state:
        st.session_state.messages_rag = []
    if 'messages_chat' not in st.session_state:
        st.session_state.messages_chat = []
    
    if 'uploaded_files_info' not in st.session_state:
        st.session_state.uploaded_files_info = []
    if 'rag_pipeline' not in st.session_state:
        st.session_state.rag_pipeline = None
    if 'index_ready' not in st.session_state:
        st.session_state.index_ready = False
    
    if 'chat_pipeline' not in st.session_state:
        st.session_state.chat_pipeline = None
    
    # ✅ AGGIUNGI: inizializzazione LLM condiviso
    if 'shared_llm' not in st.session_state:
        st.session_state.shared_llm = None
    
    # 🆕 AGGIUNGI: lista documenti accumulati
    if 'all_documents' not in st.session_state:
        st.session_state.all_documents = []

def reset_mode_data(mode):
    """Resetta i dati di una specifica modalità"""
    if mode == 'rag':
        upload_dir = Path("data/uploads")
        if upload_dir.exists():
            for file_info in st.session_state.uploaded_files_info:
                file_path = Path(file_info['file_path'])
                if file_path.exists():
                    file_path.unlink()
        
        st.session_state.uploaded_files_info = []
        st.session_state.messages_rag = []
        st.session_state.rag_pipeline = None
        st.session_state.index_ready = False
        st.session_state.all_documents = []  # 🆕 Reset documenti
    
    elif mode == 'chat':
        st.session_state.messages_chat = []
        st.session_state.chat_pipeline = None

def initialize_components():
    """Inizializza i componenti principali in base alla modalità"""
    try:
        # ✅ CREA CONFIG UNA VOLTA ALL'INIZIO
        config = RAGConfig(
            llm_model="llama3.2:3b-instruct-q4_K_M",
            embedding_model="nomic-ai/nomic-embed-text-v1.5",
            chunk_sizes=[1024, 512],
            temperature=0.3,
            context_window=4096,
            use_reranker=True,
            use_automerging=True,
            chunk_overlap=150
        )
        
        # ✅ INIZIALIZZA LLM UNA SOLA VOLTA USANDO CONFIG
        if st.session_state.shared_llm is None:
            with st.spinner("🔧 Inizializzazione LLM condiviso..."):
                from llama_index.llms.ollama import Ollama
                
                st.session_state.shared_llm = Ollama(
                    model=config.llm_model,
                    base_url=config.ollama_base_url,
                    temperature=config.temperature,
                    context_window=config.context_window,
                    request_timeout=240.0,
                )
                
                st.info(f"🔗 Connesso a Ollama: {config.ollama_base_url}")
                st.success("✅ LLM condiviso inizializzato!")
        
        if st.session_state.app_mode == 'rag':
            if st.session_state.rag_pipeline is None:
                with st.spinner("Inizializzazione RAG Pipeline..."):
                    st.session_state.rag_pipeline = OptimizedRAGPipeline(
                        config, 
                        llm=st.session_state.shared_llm
                    )
        else:
            if st.session_state.chat_pipeline is None:
                with st.spinner("Inizializzazione Chat Pipeline..."):
                    chat_config = RAGConfig(
                        llm_model=config.llm_model,
                        temperature=config.temperature,
                        context_window=config.context_window
                    )
                    st.session_state.chat_pipeline = ChatPipeline(
                        chat_config,
                        llm=st.session_state.shared_llm
                    )
        
        return True
    except Exception as e:
        st.error(f"Errore nell'inizializzazione: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return False

def main():
    """Funzione principale dell'applicazione"""
    initialize_session_state()
    
    # Mode Toggle Button
    st.markdown("""
    <div class="mode-toggle">
        <h3>🔄 Seleziona Modalità</h3>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        current_mode = st.session_state.app_mode
        
        mode_col1, mode_col2 = st.columns(2)
        
        with mode_col1:
            if st.button("📚 RAG Mode", 
                        type="primary" if current_mode == "rag" else "secondary",
                        use_container_width=True):
                if current_mode != "rag":
                    reset_mode_data("chat")
                    st.session_state.app_mode = "rag"
                    st.rerun()
        
        with mode_col2:
            if st.button("💬 Chat Mode", 
                        type="primary" if current_mode == "chat" else "secondary",
                        use_container_width=True):
                if current_mode != "chat":
                    reset_mode_data("rag")
                    st.session_state.app_mode = "chat"
                    st.rerun()
    
    # Header principale
    if st.session_state.app_mode == "rag":
        st.markdown("""
        <div class="main-header">
            🏛️ RAG Prefettura - Modalità Documenti
            <br><small>Analizza i tuoi documenti con AI</small>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="main-header-chat">
            💬 RAG Prefettura - Modalità Chat
            <br><small>Assistente legale specializzato in diritto amministrativo</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Sidebar condizionale
    if st.session_state.app_mode == "rag":
        with st.sidebar:
            st.header("📁 Gestione Documenti")
            
            # 🆕 SLIDER BATCH SIZE
            with st.expander("⚙️ Impostazioni Avanzate"):
                batch_size = st.slider(
                    "Batch Size per Indicizzazione",
                    min_value=10,
                    max_value=200,
                    value=50,
                    step=10,
                    help="Numero di nodi processati per batch. Riduci se hai file molto grandi o problemi di memoria."
                )
            
            # 🆕 CHECKBOX OCR
            use_ocr = st.checkbox(
                "🔍 Forza OCR per tutti i file",
                value=False,
                help="Se attivo, usa OCR anche per PDF/DOCX (utile per documenti scansionati)"
            )
            
            st.subheader("Carica Nuovi Documenti")
            uploaded_files = st.file_uploader(
                "Trascina i file qui o clicca per selezionare",
                accept_multiple_files=True,
                type=['pdf', 'docx', 'doc', 'txt', 'jpg', 'jpeg', 'png'],
                help="Formati supportati: PDF, DOCX, TXT, JPG, PNG"
            )
            
            if st.button("📤 Carica Documenti", disabled=not uploaded_files):
                handle_file_upload(uploaded_files, force_ocr=use_ocr, batch_size=batch_size)  # 🆕 Passa batch_size
                st.rerun()
            
            if st.session_state.uploaded_files_info:
                st.subheader("📋 Documenti Caricati")
                for doc in st.session_state.uploaded_files_info:
                    # 🆕 Icone basate sul metodo di estrazione
                    icon = {
                        'pdf_text': '📄',
                        'pdf_ocr_fallback': '🔄',
                        'ocr': '🔍',
                        'docx': '📝',
                        'text': '📃'
                    }.get(doc.get('extraction_method', 'text'), '📄')
                    
                    with st.expander(f"{icon} {doc['filename']}"):
                        st.write(f"**Tipo:** {doc['file_type']}")
                        if 'extraction_method' in doc:
                            st.write(f"**Estrazione:** {doc['extraction_method']}")
                        st.write(f"**Dimensione:** {doc['size']} bytes")
                        st.write(f"**Path:** {doc['file_path']}")
            
            st.divider()
            
            # 🎯 STATISTICHE SEMPLIFICATE
            st.subheader("📊 Statistiche Sistema")
            
            if st.session_state.rag_pipeline and st.session_state.get('index_ready', False):
                try:
                    stats = st.session_state.rag_pipeline.get_statistics()
                    
                    data = stats.get('data', {})
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("📄 Documenti", len(st.session_state.all_documents))
                    with col2:
                        st.metric("🔗 Nodi totali", data.get('total_nodes', 0))
                    
                    perf = stats.get('performance', {})
                    st.metric("💬 Query eseguite", perf.get('total_queries', 0))
                
                except Exception as e:
                    st.error(f"Errore statistiche: {e}")
            else:
                st.info("📊 Carica documenti per vedere le statistiche")
            
            st.divider()
            
            if st.button("🔄 Reset RAG"):
                reset_mode_data('rag')
                st.success("✅ RAG resettato!")
                st.rerun()
    
    else:
        with st.sidebar:
            st.header("💬 Modalità Chat")
            st.info("Chat specializzata in diritto amministrativo italiano e illeciti amministrativi.")
            
            st.subheader("📊 Statistiche")
            
            messages_count = len(st.session_state.messages_chat)
            st.metric("💬 Messaggi", messages_count)
            
            if st.session_state.chat_pipeline:
                try:
                    stats = st.session_state.chat_pipeline.get_statistics()
                    perf = stats.get('performance', {})
                    st.metric("📝 Query totali", perf.get('total_queries', 0))
                except Exception as e:
                    st.warning(f"⚠️ Statistiche non disponibili")
            
            st.divider()
            
            if st.button("🔄 Reset Chat"):
                reset_mode_data('chat')
                st.success("✅ Chat resettata!")
                st.rerun()
    
    # Area principale
    if not initialize_components():
        st.stop()
    
    chat_title = "💬 Chat Libera" if st.session_state.app_mode == "chat" else "🤖 Assistente Documenti"
    st.header(chat_title)
    
    chat_container = st.container()
    
    with chat_container:
        display_chat_history()
    
    st.divider()
    
    # Input form
    with st.form("query_form", clear_on_submit=True):
        current_mode = st.session_state.app_mode
        
        if current_mode == "rag":
            if not st.session_state.uploaded_files_info:
                st.warning("⚠️ Carica almeno un documento prima di fare domande")
            
            user_query = st.text_area(
                "Fai una domanda sui tuoi documenti:",
                placeholder="Carica prima dei documenti, poi fai domande come: 'Riassumi il contenuto' o 'Quali sono i punti chiave?'",
                height=100,
                disabled=not st.session_state.uploaded_files_info
            )
            
            submit_disabled = not st.session_state.uploaded_files_info
            submit_text = "🚀 Analizza Documenti"
        
        else:
            user_query = st.text_area(
                "Fai una domanda su diritto amministrativo:",
                placeholder="Es: 'Cosa sono gli illeciti amministrativi?' o 'Come funziona un ricorso amministrativo?'",
                height=100
            )
            
            submit_disabled = False
            submit_text = "💬 Invia"
        
        submitted = st.form_submit_button(
            submit_text,
            use_container_width=True,
            disabled=submit_disabled
        )
    
    if submitted and user_query:
        handle_query(user_query)
        st.rerun()
    
    # Footer
    st.divider()
    current_mode = st.session_state.app_mode
    
    if current_mode == "rag":
        info_title = "ℹ️ Informazioni - Modalità RAG"
        info_content = """
        **Come usare la modalità RAG:**
        
        1. **Carica documenti**: Usa la sidebar per caricare PDF, DOCX o immagini
        2. **Scegli OCR**: Spunta "Forza OCR" per documenti scansionati
        3. **Fai domande**: Scrivi domande specifiche sui tuoi file
        4. **Analisi dettagliata**: Ricevi risposte con fonti e metadata completi
        
        **Formati supportati:** PDF (nativi e scansionati), DOCX, TXT, JPG, PNG
        
        **Estrazione intelligente:**
        - PDF nativi → estrazione testo veloce
        - PDF scansionati → OCR automatico
        - Immagini → OCR diretto
        """
    else:
        info_title = "ℹ️ Informazioni - Modalità Chat"
        info_content = """
        **Come usare la modalità Chat:**
        
        - Assistente specializzato in **diritto amministrativo italiano**
        - Focus su **illeciti amministrativi** e procedure prefettizie
        - Nessun documento richiesto - conversazione diretta
        - Risposte basate su normativa italiana
        
        **Nota:** Questo è uno strumento informativo, non sostituisce la consulenza legale professionale.
        """
    
    with st.expander(info_title):
        st.markdown(info_content)

if __name__ == "__main__":
    main()